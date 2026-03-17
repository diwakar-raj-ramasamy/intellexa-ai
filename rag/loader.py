from __future__ import annotations

import json
import os
from typing import List

from langchain_core.documents import Document

from utils.settings import CHUNK_OVERLAP, CHUNK_SIZE


def _get_pdf_loader(path: str):
    # LangChain moved loaders between packages across versions; this keeps it resilient.
    try:
        from langchain_community.document_loaders import PyPDFLoader  # type: ignore
    except Exception:  # pragma: no cover
        from langchain.document_loaders import PyPDFLoader  # type: ignore
    return PyPDFLoader(path)


def _get_splitter():
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter  # type: ignore
    except Exception:  # pragma: no cover
        from langchain.text_splitter import RecursiveCharacterTextSplitter  # type: ignore
    return RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)


def load_and_chunk_pdf(path: str) -> List[Document]:
    loader = _get_pdf_loader(path)
    docs = loader.load()

    splitter = _get_splitter()
    chunks = splitter.split_documents(docs)

    # Ensure every chunk has stable metadata for UI/source display.
    for d in chunks:
        d.metadata = d.metadata or {}
        d.metadata.setdefault("source", path)
    return chunks


def _chunk_text(text: str, source: str, extra_metadata: dict | None = None) -> List[Document]:
    splitter = _get_splitter()
    docs = [Document(page_content=text or "", metadata={"source": source, **(extra_metadata or {})})]
    chunks = splitter.split_documents(docs)
    for d in chunks:
        d.metadata = d.metadata or {}
        d.metadata.setdefault("source", source)
    return chunks


def _load_docx_text(path: str) -> str:
    from docx import Document as DocxDocument  # python-docx

    doc = DocxDocument(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    return "\n".join(parts)


def _load_json_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        obj = json.load(f)
    # Pretty print preserves structure for better retrieval.
    return json.dumps(obj, ensure_ascii=False, indent=2)


def _load_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def load_and_chunk_file(path: str) -> List[Document]:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return load_and_chunk_pdf(path)
    if ext == ".docx":
        text = _load_docx_text(path)
        return _chunk_text(text, source=path, extra_metadata={"filetype": "docx"})
    if ext == ".json":
        text = _load_json_text(path)
        return _chunk_text(text, source=path, extra_metadata={"filetype": "json"})
    if ext in {".txt", ".md"}:
        text = _load_text(path)
        return _chunk_text(text, source=path, extra_metadata={"filetype": "text"})

    raise ValueError(f"Unsupported file type: {ext}")

